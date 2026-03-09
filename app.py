import os, uuid, time, json, feedparser, deepl, requests
from bs4 import BeautifulSoup
from datetime import datetime
from flask import Flask, render_template, jsonify, request, send_file, Response
from functools import wraps
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import RSS_FEEDS, DEEPL_API_KEY, TARGET_LANGUAGE

# ── RAG system ────────────────────────────────────────────────────────────────
RAG_ENABLED = bool(os.environ.get("DATABASE_URL") and os.environ.get("ANTHROPIC_API_KEY"))
if RAG_ENABLED:
    try:
        from rag import init_db, store_article_translations, rag_translate_paragraph, get_stats, add_term, get_terminology
        init_db()
        print("[APP] RAG system initialized ✓")
    except Exception as e:
        print(f"[APP] RAG init failed: {e}")
        RAG_ENABLED = False
else:
    print("[APP] RAG disabled — set DATABASE_URL and ANTHROPIC_API_KEY to enable")

app = Flask(__name__)
OUTPUT_DIR  = os.path.join(os.path.dirname(__file__), "output")

@app.after_request
def add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return response

# ── Basic Auth ────────────────────────────────────────────────────────────────
APP_USERNAME = os.environ.get("APP_USERNAME", "admin")
APP_PASSWORD = os.environ.get("APP_PASSWORD", "")

def check_auth(username, password):
    if not APP_PASSWORD:
        return True  # No password set, open access
    return username == APP_USERNAME and password == APP_PASSWORD

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not APP_PASSWORD:
            return f(*args, **kwargs)
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return Response(
                "Giriş gerekli", 401,
                {"WWW-Authenticate": 'Basic realm="Pressflow"'}
            )
        return f(*args, **kwargs)
    return decorated

FEEDS_FILE  = os.path.join(os.path.dirname(__file__), "feeds.json")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_DIR  = os.path.join(os.path.dirname(__file__), "output")
FEEDS_FILE  = os.path.join(os.path.dirname(__file__), "feeds.json")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Persistent feed storage (PostgreSQL) ─────────────────────────────────────

def get_db():
    db_url = os.environ.get("DATABASE_URL")
    if not db_url:
        return None
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(db_url, cursor_factory=psycopg2.extras.RealDictCursor)
        return conn
    except Exception as e:
        print(f"[DB] Connection error: {e}")
        return None

def init_feeds_table():
    conn = get_db()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS feeds (
                id      SERIAL PRIMARY KEY,
                name    TEXT NOT NULL,
                url     TEXT UNIQUE NOT NULL,
                enabled BOOLEAN DEFAULT TRUE,
                builtin BOOLEAN DEFAULT FALSE,
                added_at TIMESTAMP DEFAULT NOW()
            );
        """)
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"[DB] init_feeds_table error: {e}")

def load_feeds():
    """Load feeds from PostgreSQL. Falls back to JSON then config."""
    conn = get_db()
    if conn:
        try:
            cur = conn.cursor()
            init_feeds_table()
            cur.execute("SELECT name, url, enabled, builtin FROM feeds ORDER BY id")
            rows = cur.fetchall()
            cur.close()
            conn.close()
            if rows:
                return [dict(r) for r in rows]
            # First run: seed from config
            feeds = [{"name": s["name"], "url": s["url"], "enabled": True, "builtin": True}
                     for s in RSS_FEEDS]
            save_feeds(feeds)
            return feeds
        except Exception as e:
            print(f"[DB] load_feeds error: {e}")

    # Fallback: JSON
    if os.path.exists(FEEDS_FILE):
        with open(FEEDS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    feeds = [{"name": s["name"], "url": s["url"], "enabled": True, "builtin": True}
             for s in RSS_FEEDS]
    save_feeds(feeds)
    return feeds

def save_feeds(feeds):
    """Save feeds to PostgreSQL. Falls back to JSON."""
    conn = get_db()
    if conn:
        try:
            cur = conn.cursor()
            init_feeds_table()
            for feed in feeds:
                cur.execute("""
                    INSERT INTO feeds (name, url, enabled, builtin)
                    VALUES (%s, %s, %s, %s)
                    ON CONFLICT (url) DO UPDATE
                    SET name=EXCLUDED.name, enabled=EXCLUDED.enabled
                """, (feed["name"], feed["url"], feed.get("enabled", True), feed.get("builtin", False)))
            conn.commit()
            cur.close()
            conn.close()
            return
        except Exception as e:
            print(f"[DB] save_feeds error: {e}")

    # Fallback: JSON
    with open(FEEDS_FILE, "w", encoding="utf-8") as f:
        json.dump(feeds, f, ensure_ascii=False, indent=2)

def delete_feed_from_db(url):
    """Delete a feed from PostgreSQL by URL."""
    conn = get_db()
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("DELETE FROM feeds WHERE url = %s", (url,))
            conn.commit()
            cur.close()
            conn.close()
            return True
        except Exception as e:
            print(f"[DB] delete_feed error: {e}")
    return False

# ── Article helpers ──────────────────────────────────────────────────────────

def fetch_feed(source):
    try:
        feed = feedparser.parse(source["url"])
        articles = []
        for entry in feed.entries[:25]:
            articles.append({
                "id":      str(uuid.uuid4()),
                "source":  source["name"],
                "title":   entry.get("title", "No title").strip(),
                "url":     entry.get("link", ""),
                "summary": entry.get("summary", ""),
                "date":    entry.get("published", ""),
                "author":  entry.get("author", ""),
            })
        return articles
    except Exception as e:
        print(f"[ERROR] {source['name']}: {e}")
        return []

def extract_full_text(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; PressFlow/1.0)"}
        r = requests.get(url, timeout=10, headers=headers)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "lxml")
        # Remove noise
        for tag in soup(["script","style","nav","header","footer","aside","figure","iframe","noscript"]):
            tag.decompose()
        # Try article body first
        article = soup.find("article") or soup.find(class_=lambda c: c and any(x in str(c).lower() for x in ["article","content","body","story","post"]))
        target = article if article else soup.find("body")
        if not target:
            return ""
        paragraphs = [p.get_text(" ", strip=True) for p in target.find_all("p") if len(p.get_text(strip=True)) > 40]
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"[EXTRACT] {url}: {e}")
        return ""

def translate_paragraphs(translator, text, source="", author=""):
    if not text or not text.strip():
        return []
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    result = []
    for para in paragraphs:
        try:
            # Step 1: DeepL base translation
            deepl_tr = translator.translate_text(para, target_lang=TARGET_LANGUAGE).text
            # Step 2: RAG improvement (if enabled)
            if RAG_ENABLED:
                final_tr = rag_translate_paragraph(para, source=source, author=author, deepl_tr=deepl_tr)
            else:
                final_tr = deepl_tr
            result.append({"original": para, "translated": final_tr})
            time.sleep(0.05)
        except:
            result.append({"original": para, "translated": para})
    return result

# ── DOCX builder ─────────────────────────────────────────────────────────────

def build_docx(articles_data):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1.2)

    # Cover
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("HAFTALIK HABER ÖZETİ")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1A,0x3A,0x6B)
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(datetime.now().strftime("%d %B %Y"))
    dr.font.size = Pt(11); dr.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_paragraph()

    for i, art in enumerate(articles_data, 1):
        # Meta line: source + author
        meta = doc.add_paragraph()
        author_str = f"  ·  {art.get('author','')}" if art.get('author') else ""
        sr = meta.add_run(f"[{i}]  {art.get('source','').upper()}{author_str}")
        sr.bold = True; sr.font.size = Pt(9); sr.font.color.rgb = RGBColor(0x88,0x88,0x88)

        # Translated title
        titlep = doc.add_paragraph()
        tr_r = titlep.add_run(art.get("title_tr") or art.get("title",""))
        tr_r.bold = True; tr_r.font.size = Pt(14); tr_r.font.color.rgb = RGBColor(0x1A,0x3A,0x6B)

        # Original title
        op = doc.add_paragraph()
        or_r = op.add_run(f"Orijinal: {art.get('title','')}")
        or_r.italic = True; or_r.font.size = Pt(9); or_r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

        # Body
        for chunk in art.get("paragraphs", []):
            bp = doc.add_paragraph(chunk.get("translated",""))
            for r in bp.runs: r.font.size = Pt(11)

        # Source URL
        if art.get("url"):
            up = doc.add_paragraph()
            ur = up.add_run(f"Kaynak: {art['url']}")
            ur.font.size = Pt(8); ur.font.color.rgb = RGBColor(0x00,0x70,0xC0)

        # Author line (if present)
        if art.get("author"):
            authp = doc.add_paragraph()
            authr = authp.add_run(f"Yazar: {art['author']}")
            authr.font.size = Pt(9); authr.font.color.rgb = RGBColor(0x66,0x66,0x88)

        div = doc.add_paragraph("─" * 80)
        div.runs[0].font.size = Pt(8); div.runs[0].font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"Otomatik oluşturuldu • {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    filename = f"haberler_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))
    return filename

def build_text_docx(title, source, author, paragraphs):
    """Build docx for a manually pasted + translated text."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1.2)

    # Header meta
    meta = doc.add_paragraph()
    meta_str = []
    if source: meta_str.append(f"Kaynak: {source}")
    if author: meta_str.append(f"Yazar: {author}")
    mr = meta.add_run("  ·  ".join(meta_str))
    mr.font.size = Pt(9); mr.font.color.rgb = RGBColor(0x88,0x88,0x88)

    if title:
        tp = doc.add_paragraph()
        tr = tp.add_run(title)
        tr.bold = True; tr.font.size = Pt(16); tr.font.color.rgb = RGBColor(0x1A,0x3A,0x6B)

    doc.add_paragraph()

    for chunk in paragraphs:
        tp2 = doc.add_paragraph(chunk.get("translated",""))
        for r in tp2.runs: r.font.size = Pt(11)
        doc.add_paragraph()

    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    dr.font.size = Pt(8); dr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    filename = f"metin_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))
    return filename

# ── BLEU Score helper ────────────────────────────────────────────────────────

def compute_bleu(reference, hypothesis):
    """Simple BLEU-1 score between two strings."""
    import re
    def tokenize(s):
        return re.findall(r'\w+', s.lower())
    ref_tokens  = set(tokenize(reference))
    hyp_tokens  = tokenize(hypothesis)
    if not hyp_tokens:
        return 0.0
    matches = sum(1 for t in hyp_tokens if t in ref_tokens)
    return round(matches / len(hyp_tokens) * 100, 1)

# ── Routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/feeds", methods=["GET"])
def api_get_feeds():
    return jsonify({"feeds": load_feeds()})

@app.route("/api/feeds/add", methods=["POST"])
@requires_auth
def api_add_feed():
    data = request.json
    name = data.get("name","").strip()
    url  = data.get("url","").strip()
    if not name or not url:
        return jsonify({"error": "Ad ve URL gerekli"}), 400
    feeds = load_feeds()
    if any(f["name"] == name for f in feeds):
        return jsonify({"error": "Bu isimde kaynak zaten var"}), 400
    feeds.append({"name": name, "url": url, "enabled": True, "builtin": False})
    save_feeds(feeds)
    return jsonify({"ok": True})

@app.route("/api/feeds/toggle", methods=["POST"])
def api_toggle_feed():
    name = request.json.get("name","")
    feeds = load_feeds()
    for f in feeds:
        if f["name"] == name:
            f["enabled"] = not f.get("enabled", True)
            break
    save_feeds(feeds)
    return jsonify({"ok": True, "feeds": feeds})

@app.route("/api/feeds/delete", methods=["POST"])
@requires_auth
def api_delete_feed():
    url = request.json.get("url","")
    name = request.json.get("name","")
    # Delete from PostgreSQL
    if url:
        delete_feed_from_db(url)
    # Also update in-memory list
    feeds = load_feeds()
    feeds = [f for f in feeds if f["name"] != name and f.get("url","") != url]
    save_feeds(feeds)
    return jsonify({"ok": True})

@app.route("/api/fetch", methods=["POST"])
@requires_auth
def api_fetch():
    feeds = [f for f in load_feeds() if f.get("enabled", True)]
    all_articles = []
    for source in feeds:
        all_articles.extend(fetch_feed(source))
    return jsonify({"articles": all_articles, "count": len(all_articles)})

@app.route("/api/translate", methods=["POST"])
@requires_auth
def api_translate():
    data     = request.json
    selected = data.get("articles", [])
    if not selected:
        return jsonify({"error": "Makale seçilmedi"}), 400
    if DEEPL_API_KEY == "YOUR_DEEPL_API_KEY_HERE":
        return jsonify({"error": "config.py dosyasına DeepL API anahtarını ekleyin"}), 400
    try:
        translator = deepl.Translator(DEEPL_API_KEY)
    except Exception as e:
        return jsonify({"error": f"DeepL hatası: {e}"}), 500

    results = []
    for art in selected:
        body = extract_full_text(art["url"]) if art.get("url") else ""
        if not body: body = art.get("summary","")
        try: title_tr = translator.translate_text(art["title"], target_lang=TARGET_LANGUAGE).text
        except: title_tr = art["title"]
        paragraphs = translate_paragraphs(translator, body, source=art.get("source",""), author=art.get("author",""))
        article_result = {**art, "title_tr": title_tr, "paragraphs": paragraphs}
        results.append(article_result)
        # Store in RAG database
        if RAG_ENABLED:
            try: store_article_translations(article_result)
            except Exception as e: print(f"[RAG] Store error: {e}")

    filename = build_docx(results)
    return jsonify({"filename": filename, "articles": results})

@app.route("/api/translate-text", methods=["POST"])
@requires_auth
def api_translate_text():
    """Translate a manually pasted block of text."""
    data   = request.json
    text   = data.get("text","").strip()
    source = data.get("source","").strip()
    author = data.get("author","").strip()
    title  = data.get("title","").strip()

    if not text:
        return jsonify({"error": "Metin boş olamaz"}), 400
    if DEEPL_API_KEY == "YOUR_DEEPL_API_KEY_HERE":
        return jsonify({"error": "config.py dosyasına DeepL API anahtarını ekleyin"}), 400
    try:
        translator = deepl.Translator(DEEPL_API_KEY)
    except Exception as e:
        return jsonify({"error": f"DeepL hatası: {e}"}), 500

    paragraphs = translate_paragraphs(translator, text)
    title_tr   = ""
    if title:
        try: title_tr = translator.translate_text(title, target_lang=TARGET_LANGUAGE).text
        except: title_tr = title

    filename = build_text_docx(title_tr or title, source, author, paragraphs)
    return jsonify({
        "filename":   filename,
        "title_tr":   title_tr,
        "paragraphs": paragraphs,
        "source":     source,
        "author":     author,
    })

@app.route("/api/article", methods=["POST"])
def api_article():
    """Fetch full article text for detail view."""
    url = request.json.get("url", "")
    if not url:
        return jsonify({"error": "URL gerekli"}), 400
    body = extract_full_text(url)
    return jsonify({"body": body})

@app.route("/api/fetch-text", methods=["POST"])
def api_fetch_text():
    url = request.json.get("url","").strip()
    if not url:
        return jsonify({"text":""})
    text = extract_full_text(url)
    return jsonify({"text": text})

@app.route("/api/rag/stats")
def api_rag_stats():
    if not RAG_ENABLED:
        return jsonify({"enabled": False})
    return jsonify({"enabled": True, **get_stats()})

@app.route("/api/rag/terminology", methods=["GET"])
def api_get_terms():
    if not RAG_ENABLED:
        return jsonify({"terms": {}})
    return jsonify({"terms": get_terminology()})

@app.route("/api/rag/terminology", methods=["POST"])
def api_add_term():
    if not RAG_ENABLED:
        return jsonify({"error": "RAG not enabled"}), 400
    data = request.json
    term_orig = data.get("term_orig","").strip()
    term_tr   = data.get("term_tr","").strip()
    if not term_orig or not term_tr:
        return jsonify({"error": "Both fields required"}), 400
    add_term(term_orig, term_tr, data.get("source",""))
    return jsonify({"ok": True})


@app.route("/api/analytics")
def api_analytics():
    """Evaluation dashboard — works with DATABASE_URL even without RAG."""
    db_url = os.environ.get("DATABASE_URL")
    if not db_url:
        return jsonify({
            "rag_enabled": False, "archive_total": 0,
            "translations_total": 0, "sources": [],
            "recent": [], "terminology_count": 0,
        })
    try:
        import psycopg2, psycopg2.extras
        conn = psycopg2.connect(db_url, cursor_factory=psycopg2.extras.RealDictCursor)
        cur  = conn.cursor()

        archive_total = 0
        try:
            cur.execute("SELECT COUNT(*) as c FROM news_archive")
            archive_total = cur.fetchone()["c"]
        except Exception: pass

        translations_total, sources, recent, terminology_count = 0, [], [], 0
        try:
            cur.execute("SELECT COUNT(*) as c FROM translations")
            translations_total = cur.fetchone()["c"]
            cur.execute("SELECT source, COUNT(*) as count FROM translations GROUP BY source ORDER BY count DESC LIMIT 10")
            sources = [dict(r) for r in cur.fetchall()]
            # Recent: distinct articles (not Özgür Politika archive)
            cur.execute("""
                SELECT DISTINCT ON (url) source, author, title_orig, title_tr, url, created_at
                FROM translations
                WHERE source != 'Özgür Politika'
                ORDER BY url, created_at DESC
                LIMIT 10
            """)
            recent = [dict(r) for r in cur.fetchall()]
            recent.sort(key=lambda x: x.get("created_at") or "", reverse=True)
            for r in recent:
                if r.get("created_at"):
                    r["created_at"] = r["created_at"].strftime("%d.%m.%Y %H:%M")
        except Exception: pass

        try:
            cur.execute("SELECT COUNT(*) as c FROM terminology")
            terminology_count = cur.fetchone()["c"]
        except Exception: pass

        # RAG metrics
        rag_hit_rate, rag_avg_similarity, rag_total, rag_hits = 0, 0, 0, 0
        try:
            cur.execute("SELECT COUNT(*) as c FROM rag_metrics")
            rag_total = cur.fetchone()["c"]
            cur.execute("SELECT COUNT(*) as c FROM rag_metrics WHERE hit=true")
            rag_hits = cur.fetchone()["c"]
            if rag_total > 0:
                rag_hit_rate = round(rag_hits / rag_total * 100, 1)
            cur.execute("SELECT AVG(max_similarity) as s FROM rag_metrics WHERE results_found > 0")
            row = cur.fetchone()
            rag_avg_similarity = round((row["s"] or 0) * 100, 1)
        except Exception: pass

        cur.close(); conn.close()
        return jsonify({
            "rag_enabled": RAG_ENABLED,
            "archive_total": archive_total,
            "translations_total": translations_total,
            "sources": sources,
            "recent": recent,
            "terminology_count": terminology_count,
            "rag_hit_rate": rag_hit_rate,
            "rag_avg_similarity": rag_avg_similarity,
            "rag_total_queries": rag_total,
            "rag_total_hits": rag_hits,
        })
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/api/download/<filename>")
@requires_auth
def api_download(filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({"error": "Dosya bulunamadı"}), 404
    return send_file(filepath, as_attachment=True, download_name=filename)

if __name__ == "__main__":
    app.run(debug=True, port=5000)


@app.route("/api/db-status")
def db_status():
    import os, json, psycopg2, psycopg2.extras
    try:
        conn = psycopg2.connect(os.environ.get("DATABASE_URL",""), cursor_factory=psycopg2.extras.RealDictCursor)
        cur = conn.cursor()
        out = {}
        cur.execute("SELECT COUNT(*) as c FROM translations")
        out["total_paragraphs"] = cur.fetchone()["c"]
        cur.execute("SELECT COUNT(*) as c FROM translations WHERE embedding IS NOT NULL")
        out["with_embeddings"] = cur.fetchone()["c"]
        cur.execute("SELECT COUNT(DISTINCT url) as c FROM translations")
        out["unique_urls"] = cur.fetchone()["c"]
        cur.execute("""
            SELECT tc.constraint_name, string_agg(ccu.column_name, ', ' ORDER BY ccu.column_name) as cols
            FROM information_schema.table_constraints tc
            JOIN information_schema.constraint_column_usage ccu USING (constraint_name, table_name)
            WHERE tc.table_name = 'translations' AND tc.constraint_type = 'UNIQUE'
            GROUP BY tc.constraint_name
        """)
        out["unique_constraints"] = [dict(r) for r in cur.fetchall()]
        cur.execute("SELECT source, COUNT(*) as c, COUNT(DISTINCT url) as urls FROM translations GROUP BY source ORDER BY c DESC")
        out["by_source"] = [dict(r) for r in cur.fetchall()]
        cur.execute("""
            SELECT COUNT(*) as c FROM ozgurpolitika_archive a
            WHERE NOT EXISTS (SELECT 1 FROM translations t WHERE t.orig_para = a.paragraph)
        """)
        out["archive_not_yet_migrated"] = cur.fetchone()["c"]
        cur.close(); conn.close()
        return app.response_class(
            response=json.dumps(out, ensure_ascii=False, indent=2, default=str),
            mimetype="application/json"
        )
    except Exception as e:
        return app.response_class(
            response=json.dumps({"error": str(e)}, indent=2),
            mimetype="application/json"
        ), 500


@app.route("/api/fix-constraint")
def fix_constraint():
    import os, psycopg2
    try:
        conn = psycopg2.connect(os.environ.get("DATABASE_URL",""))
        conn.autocommit = True
        cur = conn.cursor()
        cur.execute("""
            SELECT constraint_name FROM information_schema.table_constraints
            WHERE table_name='translations' AND constraint_type='UNIQUE'
        """)
        dropped = []
        for row in cur.fetchall():
            c = row[0]
            if 'url_para' not in c:
                cur.execute(f"ALTER TABLE translations DROP CONSTRAINT IF EXISTS {c};")
                dropped.append(c)
        try:
            cur.execute("ALTER TABLE translations ADD CONSTRAINT translations_url_para_key UNIQUE (url, orig_para);")
        except Exception:
            pass
        cur.close(); conn.close()
        return app.response_class(
            response=json.dumps({"dropped": dropped, "status": "ok"}, indent=2),
            mimetype="application/json"
        )
    except Exception as e:
        return app.response_class(
            response=json.dumps({"error": str(e)}, indent=2),
            mimetype="application/json"
        ), 500

